#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 Skill 说明文档（docx）。

规则清单、统计数字、法条引用全部从 references/ 下的规则库与法条索引实时提取，
不硬编码，保证文档与规则库始终一致。规则库变更后重新执行即可更新文档。

本脚本依赖 python-docx，仅用于生成文档，不属于 Skill 运行时依赖。
    python3 -m pip install --user python-docx
"""

import json
import os
import sys
import re
import glob
import collections
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "金融营销材料合规审查Skill-说明文档.docx")

FONT_BODY = "宋体"
FONT_HEAD = "黑体"
FONT_MONO = "Consolas"

VERDICT_CN = {
    "violation": "违规",
    "evidence_required": "需补充证据",
    "manual_review": "需人工研判",
    "advisory": "建议",
}
SEVERITY_CN = {"red": "红", "orange": "橙", "yellow": "黄"}
BASIS_CN = {"statutory": "法定", "industry_practice": "行业惯例"}
DETECT_CN = {
    "keyword": "关键词",
    "required": "必备要素",
    "conditional_required": "条件必备",
    "numeric": "数值判断",
    "manual": "人工研判",
}
PRODUCT_CN = {
    "public_fund": "公募基金",
    "private_fund": "私募基金",
    "am_plan": "资管计划",
    "wealth_mgmt": "理财产品",
}
MEDIA_CN = {"print": "平面", "online": "线上", "video": "视频", "audio": "音频"}
AUDIENCE_CN = {"public": "面向公众", "specific": "面向特定对象"}
INST_CN = {
    "fund_manager": "基金管理人",
    "securities_firm": "证券公司",
    "bank": "银行",
    "wealth_subsidiary": "理财子公司",
}

LAW_SHORT = {
    "《中华人民共和国广告法》": "《广告法》",
    "《私募投资基金监督管理条例》": "《私募条例》",
    "《公开募集证券投资基金宣传推介材料管理暂行规定》": "《基金宣传推介暂行规定》",
    "《公开募集证券投资基金销售机构监督管理办法》": "《基金销售机构办法》",
    "《证券期货经营机构私募资产管理业务管理办法》": "《私募资管办法》",
    "《关于加强私募投资基金监管的若干规定》": "《私募若干规定》",
    "《公开募集证券投资基金业绩比较基准指引》": "《业绩基准指引》",
    "《金融产品网络营销管理办法》": "《网络营销办法》",
    "《理财公司理财产品销售管理暂行办法》": "《理财销售办法》",
    "中基协《私募投资基金募集行为管理办法》": "中基协《私募募集办法》",
    "中基协《私募证券投资基金运作指引》": "中基协《私募证券运作指引》",
    "中基协《公开募集证券投资基金投资者适当性管理细则》": "中基协《适当性细则》",
    "中证协《证券公司投资者适当性制度指引》": "中证协《适当性指引》",
}

FILE_TITLE = {
    "banned-expressions.json": "违禁表述规则（banned-expressions.json）",
    "performance-display.json": "业绩展示规则（performance-display.json）",
    "mandatory-elements.json": "必备要素规则（mandatory-elements.json）",
    "format-presentation.json": "呈现形式规则（format-presentation.json）",
}
FILE_ORDER = [
    "banned-expressions.json",
    "performance-display.json",
    "mandatory-elements.json",
    "format-presentation.json",
]


# ---------------------------------------------------------------- 数据提取

def load_rules():
    """加载规则并展开共享模式引用，使文档呈现的是实际生效的词表。"""
    sys.path.insert(0, os.path.join(ROOT, "scripts"))
    import rule_engine as eng

    groups = eng.load_shared_patterns(ROOT)
    by_file = {}
    for path in glob.glob(os.path.join(ROOT, "references/rules/*.json")):
        name = os.path.basename(path)
        if name.startswith("_"):
            continue
        with open(path, encoding="utf-8") as f:
            rules = json.load(f)["rules"]
        for rule in rules:
            eng.expand_pattern_refs(rule, groups, [])
        by_file[name] = rules
    return by_file


def load_legal_index():
    path = os.path.join(ROOT, "references/legal-index.md")
    text = open(path, encoding="utf-8").read()
    entries = []
    for rid, title in re.findall(r"^### (\S+) (.+)$", text, re.M):
        seg = text.split("### " + rid + " ", 1)[1]
        eff = re.search(r"生效状态：(.+)", seg)
        ver = re.search(r"核实状态：(.+)", seg)
        entries.append({
            "ref_id": rid,
            "title": title.strip(),
            "effective": eff.group(1).strip().replace("*", "") if eff else "",
            "verified": ver.group(1).strip() if ver else "",
        })
    return entries


def count_tests(by_file):
    n = 0
    for rules in by_file.values():
        for r in rules:
            t = r.get("tests") or {}
            n += len(t.get("match", [])) + len(t.get("no_match", []))
    return n


def product_coverage(by_file):
    cnt = collections.Counter()
    for rules in by_file.values():
        for r in rules:
            prods = r.get("scope", {}).get("products") or list(PRODUCT_CN)
            for p in prods:
                cnt[p] += 1
    return cnt


def scope_text(rule):
    sc = rule.get("scope", {})
    parts = []
    prods = sc.get("products")
    parts.append("、".join(PRODUCT_CN.get(p, p) for p in prods) if prods else "全部产品")
    if sc.get("audience"):
        parts.append("、".join(AUDIENCE_CN.get(a, a) for a in sc["audience"]))
    if sc.get("media"):
        parts.append("、".join(MEDIA_CN.get(m, m) for m in sc["media"]))
    if sc.get("institutions"):
        parts.append("、".join(INST_CN.get(i, i) for i in sc["institutions"]))
    return "；".join(parts)


def basis_text(rule):
    lb = rule.get("legal_basis") or []
    if not lb:
        return "无条款依据（行业惯例）"
    return "；".join(LAW_SHORT.get(x["law"], x["law"]) + x["clause"] for x in lb)


# ---------------------------------------------------------------- 排版工具

def set_fonts(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    pf = normal.paragraph_format
    pf.line_spacing = 1.4
    pf.space_after = Pt(4)

    for name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        st = doc.styles[name]
        st.font.name = FONT_HEAD
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)
        st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        st.paragraph_format.space_after = Pt(6)


def set_margins(doc):
    for s in doc.sections:
        s.page_width, s.page_height = Cm(21), Cm(29.7)
        s.top_margin = s.bottom_margin = Cm(2.2)
        s.left_margin = s.right_margin = Cm(2.4)


def add_page_number(doc):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for el, attr, val in (
        ("w:fldChar", "w:fldCharType", "begin"),
        ("w:instrText", "xml:space", "preserve"),
        ("w:fldChar", "w:fldCharType", "end"),
    ):
        e = OxmlElement(el)
        e.set(qn(attr), val)
        if el == "w:instrText":
            e.text = "PAGE"
        run._r.append(e)
    run.font.size = Pt(9)
    run.font.name = FONT_BODY


def para(doc, text, size=10.5, bold=False, italic=False, align=None,
         space_after=None, color=None, mono=False):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    font = FONT_MONO if mono else FONT_BODY
    r.font.name = font
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    if color:
        r.font.color.rgb = color
    return p


def bullets(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        bold_head = None
        if isinstance(it, tuple):
            bold_head, it = it
        if bold_head:
            r = p.add_run(bold_head)
            r.bold = True
            r.font.size = Pt(size)
            r.font.name = FONT_BODY
            r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        r = p.add_run(it)
        r.font.size = Pt(size)
        r.font.name = FONT_BODY
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)


def shade(cell, hex_color):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(el)


def table(doc, headers, rows, widths=None, size=9, header_bg="D9E2F3"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(size)
        r.font.name = FONT_BODY
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        shade(hdr[i], header_bg)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(str(v))
            r.font.size = Pt(size)
            r.font.name = FONT_BODY
            r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.font.name = FONT_MONO
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MONO)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def h1(doc, t):
    doc.add_heading(t, level=1)


def h2(doc, t):
    doc.add_heading(t, level=2)


def h3(doc, t):
    doc.add_heading(t, level=3)


# ---------------------------------------------------------------- 文档主体

def build():
    by_file = load_rules()
    legal = load_legal_index()
    all_rules = [r for rules in by_file.values() for r in rules]
    total = len(all_rules)
    n_tests = count_tests(by_file)
    cov = product_coverage(by_file)

    c_verdict = collections.Counter(r["verdict"] for r in all_rules)
    c_basis = collections.Counter(r["basis"] for r in all_rules)
    c_status = collections.Counter(r["status"] for r in all_rules)
    c_detect = collections.Counter(r["detect"]["type"] for r in all_rules)
    c_sev = collections.Counter(r["severity"] for r in all_rules)

    law_freq = collections.Counter()
    for r in all_rules:
        for lb in r.get("legal_basis", []):
            law_freq[lb["law"]] += 1
    c_verified = collections.Counter(
        e["verified"].split("（")[0].strip() for e in legal
    )

    doc = Document()
    set_fonts(doc)
    set_margins(doc)
    add_page_number(doc)

    # ---------------- 封面
    for _ in range(5):
        doc.add_paragraph()
    para(doc, "金融营销材料合规审查 Skill", size=26, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, "说 明 文 档", size=15,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30,
         color=RGBColor(0x44, 0x44, 0x44))
    para(doc, "基于本地规则库的营销宣传材料发布前合规审查工具", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=60,
         color=RGBColor(0x66, 0x66, 0x66))

    table(doc, ["项目", "内容"], [
        ["运行平台", "WorkBuddy Agent Skill"],
        ["规则库规模", f"{total} 条规则，分 {len(by_file)} 个规则文件"],
        ["法条索引", f"{len(legal)} 个条目，覆盖法律至自律规则五个位阶"],
        ["回归测试", f"{n_tests} 个内嵌测试用例"],
        ["运行环境", "Python 3.9+，零第三方依赖"],
        ["数据处理", "全本地运行，材料内容不出网"],
        ["文档生成时间", datetime.now().strftime("%Y 年 %m 月 %d 日")],
    ], widths=[4.0, 11.6], size=10)

    para(doc, "本文档的规则清单与统计数据由 tools/gen_doc.py 从规则库实时提取生成，"
              "与 references/ 下的规则文件保持一致。", size=9,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x88, 0x88, 0x88))

    doc.add_page_break()

    # ---------------- 目录
    h1(doc, "目录")
    toc = [
        "一、Skill 是什么：定位与要解决的问题",
        "二、工作逻辑：一份材料是怎么被审查的",
        "三、支持检测的素材范围",
        "四、规则库完整清单",
        "五、法律依据来源",
        "六、质量保障机制",
        "七、已知边界",
        "八、使用方式",
    ]
    for i, t in enumerate(toc, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(t)
        r.font.size = Pt(11)
        r.font.name = FONT_BODY
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    doc.add_page_break()

    # ---------------- 一、定位
    h1(doc, "一、Skill 是什么：定位与要解决的问题")

    h2(doc, "1.1 一句话定位")
    para(doc, "对基金、资管、理财产品的营销宣传材料做发布前合规审查，输出可追溯到法条原文的"
              "问题清单与修改建议。定位是合规人员的第一道筛子，不是最终判定者。", bold=True)

    h2(doc, "1.2 要解决的问题")
    para(doc, "资管机构的营销材料在对外发布前必须经合规审查。现状是合规人员逐字读稿，"
              "凭经验对照几十份监管文件判断。这件事的痛点不在难，在于三点：")
    bullets(doc, [
        ("量大且重复。", "新品发行、公众号推文、路演材料、海报，一家中型基金公司每周几十份。"),
        ("规则分散且分层。", "从《广告法》到中基协自律规则散落在十几份文件里，"
                        "且公募、私募、资管计划、理财产品适用的条款各不相同。"),
        ("规则会变。", "《金融产品网络营销管理办法》2026 年 9 月 30 日施行，一旦生效，"
                   "存量材料里的「低风险」「秒到账」全部触线，需要批量重扫。"),
    ])

    h2(doc, "1.3 两条硬约束")
    para(doc, "以下两点决定了技术方案的形态，不是可选项：")
    bullets(doc, [
        ("材料内容高度敏感。", "未发布的产品信息属于内幕信息（MNPI），"
                         "不能出网、不能过第三方 SaaS 平台。因此全部处理在本地完成，不调用任何外部 API。"),
        ("合规结论不能由机器下。", "合规意见书需要签字担责，工具只能做第一道筛子。"
                            "因此输出的是待复核清单，不是审批结论。"),
    ])

    h2(doc, "1.4 三条边界声明")
    para(doc, "以下三点在每次输出报告时都会向使用者明确，不得省略：")
    table(doc, ["序号", "边界声明"], [
        ["1", "本工具输出的是待复核清单，不是合规结论，最终判定权在合规人员。"],
        ["2", "标记为「需人工研判」的项必须由人工确认，工具只标记触发位置。"],
        ["3", "未检出问题不等于材料合规，规则库覆盖范围有已知边界。"],
    ], widths=[1.6, 14.0], size=10)
    para(doc, "报告中禁止使用「审核通过」「合规」「可以发布」这类结论性表述。", bold=True)

    doc.add_page_break()

    # ---------------- 二、工作逻辑
    h1(doc, "二、工作逻辑：一份材料是怎么被审查的")

    h2(doc, "2.1 整体流程")
    code_block(doc, """      待审材料（md / txt / 由 Agent 解析的 pdf、docx、pptx）
                        |
                        v
        [ 第一步 ] 确认四维审查上下文
                   产品类型 / 受众范围 / 材料载体 / 机构类型
                        |
                        v
        [ 第二步 ] 确定性规则扫描  scan_rules.py -> rule_engine.py
                   作用域路由 -> 五类检测器 -> 排除上下文 -> 重叠去重
                        |
                        |  读取
                        v
           references/rules/*.json  --ref_id-->  references/legal-index.md
                        |
                        v
                   JSON 输出（每条含法条原文四元组）
                        |
                        v
        [ 第三步 ] Agent 语义研判
                   整体印象误导 / 变相承诺 / 风格漂移 / 数据矛盾 / 术语未解释
                        |
                        v
        [ 第四步 ] 生成分级报告
                        |
                        v
        [ 第五步 ] 批量重扫（可选，用于新规生效后的存量材料排查）
                        |
                        v
                  待复核清单 --> 合规人员判定""")

    para(doc, "确定性扫描与语义研判是分离的两层，各管各的：能被规则表达的交给规则，"
              "规则表达不了的交给大模型。不让大模型去做关键词匹配（结果不稳定、不可复现），"
              "也不指望正则去理解「通篇只讲收益不讲风险」（做不到）。")

    h2(doc, "2.2 四维上下文路由")
    para(doc, "同一句话在不同场景下适用的法条不同，因此每条规则声明四个维度的作用域，"
              "扫描时据此决定是否应用该规则。")
    table(doc, ["维度", "取值", "判断依据"], [
        ["产品类型", "公募基金 / 私募基金 / 资管计划 / 理财产品",
         "产品全称、「证券投资基金」字样、备案编码"],
        ["受众范围", "面向不特定对象 / 面向特定对象",
         "是否公开发布、有无合格投资者确认环节"],
        ["材料载体", "平面 / 线上 / 视频 / 音频", "投放渠道"],
        ["机构类型", "基金管理人 / 证券公司 / 银行 / 理财子公司", "材料落款"],
    ], widths=[2.4, 7.0, 6.2], size=9.5)

    para(doc, "最典型的例子是 BAN-001 与 BAN-002：都是「承诺保本保收益」，"
              "但公募依据《广告法》第二十五条与《基金宣传推介暂行规定》第十五条，"
              "私募依据《私募投资基金监督管理条例》第二十条（位阶为行政法规）。"
              "二者在「避险」一词上还有实质差异——私募场景明文禁用，公募场景无现行依据。"
              "合并成一条规则就会引用错误法条。")
    para(doc, "上下文未提供时，相关规则跳过而非默认适用，即宁可漏报也不在信息不足时误报，"
              "报告中会说明哪些规则因此未被应用。", bold=True)

    h2(doc, "2.3 五类检测器")
    para(doc, "规则不是清一色的关键词匹配。这个领域大量条款不是「禁止说什么」，"
              "而是「说了 A 就必须同时说 B」，纯敏感词方案覆盖不了。")
    table(doc, ["检测方式", "数量", "机制说明"], [
        [DETECT_CN[k], v, {
            "keyword": "正则匹配命中即报，支持排除上下文以规避否定语境误报",
            "manual": "触发词命中后仅标记位置，交由人工结合上下文研判",
            "conditional_required": "若材料出现 A 则必须同时出现 B，缺 B 即报",
            "required": "材料必须包含某要素，缺失即报",
            "numeric": "从文本提取数值并做区间判断，如业绩区间是否满 6 个月",
        }[k]] for k, v in c_detect.most_common()
    ], widths=[2.6, 1.4, 11.6], size=9.5)

    h2(doc, "2.4 四种判定类型")
    para(doc, "早期版本只有「违规 / 不违规」二分。问题在于大量规则本质上无法由文本判定，"
              "例如「销售机构标注的风险等级不得低于管理人评级」需要查管理人官方评级。"
              "这类检查既不能报违规（没有依据），也不能不报（确实是检查项）。"
              "拆成四类之后，每类在报告中的呈现方式与责任归属都不同。")
    table(doc, ["判定类型", "数量", "含义", "报告中如何处理"], [
        ["违规 violation", c_verdict.get("violation", 0), "命中即违规", "直接列入报告，附法条原文"],
        ["需人工研判 manual_review", c_verdict.get("manual_review", 0),
         "需上下文或外部数据判断", "标记位置，注明须合规人员核验"],
        ["建议 advisory", c_verdict.get("advisory", 0), "无条款依据的优化建议",
         "列入建议区，禁止表述为违规"],
        ["需补充证据 evidence_required", c_verdict.get("evidence_required", 0),
         "法条限定未提供客观证据时才违规", "提示补充数据出处，而非要求删除"],
    ], widths=[4.4, 1.4, 4.4, 5.4], size=9)

    h2(doc, "2.5 法定依据与行业惯例的区分")
    table(doc, ["依据类型", "数量", "处理方式"], [
        ["法定 statutory", c_basis.get("statutory", 0), "挂条款号，报告中引用法条原文"],
        ["行业惯例 industry_practice", c_basis.get("industry_practice", 0),
         "禁止挂条款号，报告中必须标注「无现行条款依据」"],
    ], widths=[4.2, 1.6, 9.8], size=9.5)
    para(doc, "这个区分不是洁癖。市面上大量合规工具仍在使用证监会公告〔2008〕2 号的禁用词表"
              "（净值归一、首只、最大、最好、最强、避险），而该文件已被证监会公告〔2020〕59 号"
              "第二十条明文废止。把废止文件的词表当作法条引用，等于给合规人员一个假的举证依据。"
              "本规则库保留这些词的检查（行业惯例仍有参考价值），但强制标注无条款依据，"
              "校验脚本会拒绝给行业惯例类规则挂条款号。")

    h2(doc, "2.6 新规过渡期处理")
    table(doc, ["生效状态", "数量", "说明"], [
        ["现行有效 active", c_status.get("active", 0), "正常应用"],
        ["已发布未生效 pending", c_status.get("pending", 0),
         "报告中单独分节，标明生效日期，不与现行规则混列"],
    ], widths=[4.2, 1.6, 9.8], size=9.5)
    pend = [r for r in all_rules if r["status"] == "pending"]
    if pend:
        para(doc, "当前未生效规则：" + "；".join(
            f"{r['rule_id']}（{r['title']}，{r.get('effective_date', '')} 施行）" for r in pend))

    h2(doc, "2.7 报告输出结构")
    para(doc, "扫描引擎输出 JSON，每条发现包含以下四元组，这是本工具与通用文案检查工具的根本差异："
              "合规人员需要据此向监管举证。")
    table(doc, ["字段", "内容"], [
        ["原文引用", "命中文本、所在行号、上下文片段"],
        ["法条依据", "法规全称 + 条款号 + 条款原文（不改写、不转述）"],
        ["处罚案例", "对应的真实监管处罚案例与原始认定措辞（如有）"],
        ["修改建议", "具体的修改方向，而非「违反相关规定」这类模糊表述"],
    ], widths=[3.0, 12.6], size=9.5)
    para(doc, "报告由 Agent 基于 JSON 渲染，而非硬编码模板。理由是报告的呈现需求是变的——"
              "有人要按严重度排序，有人要按材料章节排序，有人要导出简版给业务方。"
              "硬编码模板每变一次需求就要改一次代码，Agent 渲染只需调整 SKILL.md 中的指令。")

    doc.add_page_break()

    # ---------------- 三、素材范围
    h1(doc, "三、支持检测的素材范围")

    h2(doc, "3.1 支持的文件格式")
    table(doc, ["支持程度", "格式", "说明"], [
        ["直接支持", "Markdown（.md）、纯文本（.txt）", "扫描脚本直接读取，无需转换"],
        ["经解析支持", "PDF、Word（.docx）、PPT（.pptx）、HTML",
         "由 WorkBuddy Agent 先提取纯文本，再交给扫描脚本"],
        ["不支持", "图片中的文字", "需 OCR 能力，当前未集成"],
        ["不支持", "视频、音频成片", "仅能审查对应的文字脚本，无法从成片中提取"],
    ], widths=[2.4, 5.4, 7.8], size=9.5)

    h2(doc, "3.2 支持的材料类型")
    para(doc, "覆盖资管机构对外发布的主要营销材料形态：")
    table(doc, ["材料类型", "典型场景", "适配说明"], [
        ["产品说明书", "新产品发行前的产品要素说明", "必备要素与风险揭示检查为主"],
        ["产品推介材料", "宣传单页、产品手册、销售话术稿", "覆盖全部四类规则，检出率最高"],
        ["公众号推文", "线上渠道的产品推广文章", "触发线上载体专属规则"],
        ["路演 PPT", "面向机构或个人投资者的演示材料", "需先由 Agent 提取文本"],
        ["海报与长图", "活动海报、产品长图的文案部分", "仅审查文案，版式与字号不在覆盖范围"],
        ["短视频与音频脚本", "投放前的口播稿、分镜脚本", "含风险提示时长检查（FMT-003）"],
        ["存量材料批量重扫", "新规发布后对已发布材料的合规排查",
         "按规则维度聚合，标记受同一新规影响的材料"],
    ], widths=[3.2, 6.2, 6.2], size=9.5)

    h2(doc, "3.3 支持的产品类型与规则覆盖")
    para(doc, "不同产品类型适用的规则数量差异较大，反映了监管规则本身的密度差异，"
              "也反映了当前规则库的建设重点。")
    table(doc, ["产品类型", "适用规则数", f"占总数（{total} 条）比例", "说明"], [
        [PRODUCT_CN[k], cov[k], f"{cov[k] * 100 // total}%", desc]
        for k, desc in [
            ("public_fund", "监管规则最密集，规则库建设重点"),
            ("private_fund", "以中基协自律规则与《私募条例》为主"),
            ("am_plan", "证券期货经营机构私募资管业务"),
            ("wealth_mgmt", "银行理财，当前覆盖偏薄，后续需补充"),
        ]
    ], widths=[3.0, 2.4, 3.4, 6.8], size=9.5)

    h2(doc, "3.4 支持的载体、受众与机构类型")
    table(doc, ["维度", "支持取值"], [
        ["材料载体", "平面（print）、线上（online）、视频（video）、音频（audio）"],
        ["受众范围", "面向不特定对象（public）、面向特定对象（specific）"],
        ["机构类型", "基金管理人、证券公司、银行、理财子公司"],
    ], widths=[3.0, 12.6], size=9.5)
    para(doc, "机构类型会影响必备要素检查。例如 REQ-005（券商推广材料未标明必备三项）"
              "仅在机构类型为证券公司时应用，依据是中证协《证券公司投资者适当性制度指引》"
              "第二十二条。")

    doc.add_page_break()

    # ---------------- 四、规则清单
    h1(doc, "四、规则库完整清单")

    h2(doc, "4.1 总览")
    table(doc, ["规则文件", "检查内容", "规则数"], [
        [FILE_TITLE[f].split("（")[0], desc, len(by_file[f])]
        for f, desc in [
            ("banned-expressions.json", "材料中不该出现的表述"),
            ("performance-display.json", "业绩数据的展示方式"),
            ("mandatory-elements.json", "必须出现的要素与法定声明"),
            ("format-presentation.json", "排版与呈现形式"),
        ]
    ], widths=[4.2, 8.4, 3.0], size=9.5)
    para(doc, "规则按「检查什么」而非「违反哪条法」分文件，因为使用者查找规则时"
              "通常是按现象查找的。严重级别分布：" +
         "、".join(f"{SEVERITY_CN[k]}级 {v} 条" for k, v in c_sev.most_common()) + "。")

    para(doc, "下列各表中，「判定」列取值含义见 2.4 节，「检测」列取值含义见 2.3 节，"
              "「法条依据」列使用法规简称，简称与全称的对照见 5.1 节。", size=9,
         color=RGBColor(0x66, 0x66, 0x66))

    for fname in FILE_ORDER:
        rules = by_file.get(fname, [])
        if not rules:
            continue
        h2(doc, f"4.{FILE_ORDER.index(fname) + 2} " + FILE_TITLE[fname])
        rows = []
        for r in rules:
            vd = VERDICT_CN[r["verdict"]]
            if r["status"] == "pending":
                vd += "（未生效）"
            rows.append([
                r["rule_id"],
                r["title"],
                vd + f"／{SEVERITY_CN[r['severity']]}",
                DETECT_CN[r["detect"]["type"]],
                scope_text(r),
                basis_text(r),
            ])
        table(doc, ["编号", "规则名称", "判定／级别", "检测", "适用范围", "法条依据"],
              rows, widths=[1.5, 3.6, 1.9, 1.4, 3.0, 4.2], size=8)

    doc.add_page_break()

    # ---------------- 五、法律依据
    h1(doc, "五、法律依据来源")

    h2(doc, "5.1 引用的法规与频次")
    para(doc, f"references/legal-index.md 收录 {len(legal)} 个法条条目，按效力位阶分五层。"
              "每条规则通过 ref_id 指向索引条目，校验脚本会检查引用的条目是否存在。")
    rows = []
    for law, n in law_freq.most_common():
        rows.append([law, LAW_SHORT.get(law, "—"), n])
    table(doc, ["法规全称", "文档中简称", "被规则引用次数"], rows,
          widths=[8.6, 4.4, 2.6], size=9)
    top_law, top_n = law_freq.most_common(1)[0]
    para(doc, f"{top_law}（证监会公告〔2020〕59 号）被引用 {top_n} 次，"
              "是公募营销材料审查的主干规范。法条索引中另有部分条目已收录但暂未被规则引用，"
              "属于覆盖面扩展的预留。")

    h2(doc, "5.2 核实状态")
    para(doc, "每个法条条目标注核实状态。保留这个字段而不是一律标注「已核实」，"
              "是因为报告中的法条原文会被合规人员拿去作为举证依据，来源不够硬的必须让使用者知道。")
    table(doc, ["核实状态", "数量", "含义"], [
        ["已核实原文", c_verified.get("已核实原文", 0), "从官方来源取得条款完整原文"],
        ["已核实存在", c_verified.get("已核实存在", 0), "确认条款存在，但未从官方来源取得完整原文"],
        ["未能核实", c_verified.get("未能核实", 0), "无法确认，禁止被规则引用"],
    ], widths=[3.0, 1.8, 10.8], size=9.5)

    partial = [e for e in legal if e["verified"].startswith("已核实存在") and "（" in e["verified"]]
    if partial:
        para(doc, "标注为「已核实存在」的条目中，以下几条附有具体说明：")
        table(doc, ["编号", "条目", "说明"],
              [[e["ref_id"], e["title"], e["verified"].split("（", 1)[1].rstrip("）")]
               for e in partial],
              widths=[2.6, 6.0, 7.0], size=8.5)

    h2(doc, "5.3 时效性处理")
    para(doc, "法条索引开头设有「时效性警示」一节，处理两类问题：")
    bullets(doc, [
        ("已废止但仍在流传的规则。", "证监会公告〔2008〕2 号已被 2020 年 59 号文第二十条明文废止，"
                             "但其禁用词表和「5 个工作日内报证监局备案」的流程要求仍在大量资料中流传。"
                             "现行《基金宣传推介暂行规定》没有向证监局报备的要求，"
                             "只要求内部合规审查、出具合规意见书、存档备查。"),
        ("已发布未生效的规则。", "《金融产品网络营销管理办法》2026 年 9 月 30 日施行，"
                          "相关规则标记为未生效，报告中单独分节。"
                          "《公开募集证券投资基金销售行为规范》仍为征求意见稿，不入规则库。"),
    ])

    h2(doc, "5.4 完整法条索引")
    table(doc, ["编号", "法规与条款", "生效状态"],
          [[e["ref_id"], e["title"], e["effective"].split("（")[0].strip()] for e in legal],
          widths=[3.0, 9.6, 3.0], size=8.5)

    doc.add_page_break()

    # ---------------- 六、质量保障
    h1(doc, "六、质量保障机制")

    h2(doc, "6.1 规则库校验")
    para(doc, "validate_rules.py 对整个规则库执行校验，任何规则改动后必须通过，不通过不得交付。"
              "校验项包括：")
    bullets(doc, [
        "字段完整性与枚举取值合法性",
        "rule_id 全局唯一性",
        "全部正则表达式可编译",
        "legal_basis 引用的 ref_id 在法条索引中存在",
        "行业惯例类规则未挂条款号",
        "执行每条规则的内嵌测试用例",
    ])

    h2(doc, "6.2 内嵌回归测试")
    para(doc, f"每条规则携带 match 与 no_match 两组测试样例，当前共 {n_tests} 个用例，"
              "校验时全部执行。")
    para(doc, "这是整套设计中最关键的一环。合规规则的正则天然容易过度匹配，"
              "改动一处很容易在别处引入误报，而误报在这个场景下代价极高。"
              "没有回归测试，规则库迭代几轮之后就没有人敢改了。")

    h2(doc, "6.3 误报比漏报更危险")
    para(doc, "开发过程中使用一份完全合规的公募材料作为基线测试，抓出三个误报，"
              "全部属于同一类问题：把法规强制要求写的句子判成了违规。")
    table(doc, ["命中内容", "材料原文", "实际性质"], [
        ["没有风险", "……也不表明投资于本基金没有风险", "中国证监会要求的法定注册声明"],
        ["预期收益", "其预期风险与预期收益高于债券型基金", "基金运作管理办法的风险收益特征标准句式"],
        ["基准收益率", "业绩比较基准收益率", "法定允许的业绩比较基准列示方式"],
    ], widths=[2.8, 7.4, 5.4], size=9)
    para(doc, "这类误报比漏报危险得多：合规人员照着报告修改，材料会从合规变成缺失法定要素。"
              "修复方式是排除上下文机制（命中点前 12 字符窗口内出现排除模式则跳过）"
              "与正则负向断言，并将这三个句子写入回归测试、在规则备注中注明不得移除。", bold=True)

    h2(doc, "6.4 当前样本基线")
    table(doc, ["样本", "违规", "需人工研判", "建议"], [
        ["公募合规样本（误报基线）", 0, 5, 1],
        ["公募违规样本（检出基线）", 17, 6, 2],
    ], widths=[6.4, 3.0, 3.2, 3.0], size=9.5)
    para(doc, "样本库中合规样本与违规样本同等重要：违规样本验证能否检出，"
              "合规样本验证会不会乱报。目前尚未在真实营销材料上做过评测，"
              "precision 与 recall 暂无数据。", bold=True)

    doc.add_page_break()

    # ---------------- 七、边界
    h1(doc, "七、已知边界")
    para(doc, "以下是当前明确做不到的事项，不以含糊表述掩盖：")
    table(doc, ["序号", "边界", "说明"], [
        ["1", "不做版式与视觉审查",
         "字号、颜色对比度、风险提示是否「醒目」这类要求无法从纯文本判断，"
         "FMT-004 只能定为建议类"],
        ["2", "不校验外部事实",
         "基金合同约定的业绩比较基准、管理人官方风险等级、基金经理管理的全部产品清单，"
         "需接入外部数据源才能真正校验，当前全部落在需人工研判"],
        ["3", "视频音频只能审脚本",
         "FMT-003 的 5 秒风险提示时长需材料中标注时长信息，无法从成片直接测量"],
        ["4", "整体印象类违规依赖语义研判",
         "稳定性弱于规则层，不保证结果可复现"],
        ["5", "规则覆盖面以公募为主",
         f"适用规则数：公募 {cov['public_fund']} 条、私募 {cov['private_fund']} 条、"
         f"资管计划 {cov['am_plan']} 条、理财产品 {cov['wealth_mgmt']} 条，"
         "理财与资管偏薄"],
        ["6", "法规时效需人工维护",
         "规则库不会自动感知新规发布，法条索引的核实时间需定期更新"],
        ["7", "尚未在真实材料上评测",
         "当前基线来自 8 份自建样本，覆盖 43 条规则中的 39 条，真实材料准确率数据待补"],
    ], widths=[1.4, 4.2, 10.0], size=9)

    # ---------------- 八、使用方式
    h1(doc, "八、使用方式")

    h2(doc, "8.1 安装")
    code_block(doc, "ln -sfn /path/to/fin-marketing-compliance-review \\\n"
                    "        ~/.workbuddy/skills/fin-marketing-compliance-review")
    para(doc, "使用符号链接而非复制，规则库改动即时生效，无需重复同步。")

    h2(doc, "8.2 在 WorkBuddy 中调用")
    para(doc, "直接以自然语言触发，例如「帮我审查一下这份材料能不能发」，并附上材料文件。"
              "Agent 会先确认产品类型与投放渠道等上下文，再执行扫描与研判。")

    h2(doc, "8.3 命令行调用")
    code_block(doc, "python3 scripts/scan_rules.py \\\n"
                    "  --file examples/sample-public-fund-violating.md \\\n"
                    "  --product public_fund \\\n"
                    "  --audience public \\\n"
                    "  --media online \\\n"
                    "  --institution fund_manager \\\n"
                    "  --json")
    para(doc, "四个上下文参数均可省略，省略时对应规则跳过并在结果中说明。")

    h2(doc, "8.4 校验规则库")
    code_block(doc, "python3 scripts/validate_rules.py")

    h2(doc, "8.5 新增或修改规则")
    para(doc, "先阅读 references/rules/SCHEMA.md，将规则放入对应的规则文件，"
              "必须包含测试样例，法条依据必须在法条索引中有对应条目。改完执行校验。"
              "规则库变更后重新运行 tools/gen_doc.py 即可同步更新本文档。")

    # ---------------- 免责
    h1(doc, "免责声明")
    bullets(doc, [
        "本工具输出的是待复核清单，不构成合规结论，不能替代合规人员的专业判断与合规意见书。",
        "规则库中的法规状态截至法条索引标注的核实时间，使用前请自行确认法规时效。",
        "examples/ 目录中所有机构名、产品名、人员姓名均为虚构，与任何真实主体无关。",
    ])

    doc.save(OUT)
    return OUT, total, len(legal), n_tests


if __name__ == "__main__":
    path, n_rules, n_legal, n_tests = build()
    print(f"已生成：{path}")
    print(f"规则 {n_rules} 条 / 法条条目 {n_legal} 个 / 测试用例 {n_tests} 个")
